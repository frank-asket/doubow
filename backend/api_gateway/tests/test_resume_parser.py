import io

import pytest
from docx import Document

from services.resume_parser import (
    extract_plain_text,
    heuristic_profile_from_text,
)


def test_heuristic_profile_extracts_skills_years_headline():
    text = """Jordan Lee
Senior Machine Learning Engineer

8 years of experience shipping models with Python, PyTorch, and AWS.

Skills: Kubernetes, PostgreSQL, GraphQL
"""
    d = heuristic_profile_from_text(text)
    assert "Senior Machine Learning Engineer" == d["headline"] or "Senior" in d["headline"]
    assert "Python" in d["skills"]
    assert "PyTorch" in d["skills"]
    assert d["experience_years"] >= 8
    assert "Jordan Lee" == d["name"]


def test_heuristic_prefers_long_skill_phrases():
    lower = """Experience with Natural Language Processing and Python."""
    d = heuristic_profile_from_text(lower)
    assert "Natural Language Processing" in d["skills"]
    assert "Python" in d["skills"]


@pytest.mark.asyncio
async def test_docx_extract_and_profile():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Taylor Morgan")
    doc.add_paragraph("Staff Platform Engineer — Go, Kafka, Redis")
    doc.add_paragraph("Skills: Terraform, Docker, CI/CD across 6 years.")
    doc.save(buf)
    raw = buf.getvalue()

    plain = extract_plain_text(raw, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Taylor Morgan" in plain
    assert "Kafka" in plain

    profile = heuristic_profile_from_text(plain)
    assert profile["headline"]
    assert "Kafka" in profile["skills"]
    assert profile["experience_years"] >= 6
