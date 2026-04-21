"""Extract plain text from résumé uploads and derive a minimal structured profile."""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from pypdf import PdfReader

# Ordered longest-first for skill deduplication when scanning.
_SKILL_PHRASES: tuple[str, ...] = (
    "Natural Language Processing",
    "Machine Learning",
    "Deep Learning",
    "Large Language Model",
    "Computer Vision",
    "Distributed Systems",
    "Site Reliability Engineering",
    "Ruby on Rails",
    "Spring Boot",
    "PostgreSQL",
    "JavaScript",
    "TypeScript",
    "Elasticsearch",
    "Kubernetes",
    "Terraform",
    "FastAPI",
    "Next.js",
    "Node.js",
    "GraphQL",
    "MongoDB",
    "PyTorch",
    "TensorFlow",
    "LangChain",
    "Angular",
    "Kotlin",
    "Scala",
    "Python",
    "Java",
    "Rust",
    "Ruby",
    "Swift",
    "React",
    "Vue.js",
    "Docker",
    "pandas",
    "NumPy",
    "Redis",
    "AWS",
    "GCP",
    "Azure",
    "Django",
    "Flask",
    "Apache Spark",
    "Kafka",
    "Git",
    "SQL",
    "RAG",
    "LLM",
    "MLOps",
    "CI/CD",
    "Go",
    "C++",
)

_YEARS_PATTERNS = (
    re.compile(r"(?:over\s+)?(\d{1,2})\+?\s*years?\s+of\s+(?:experience|professional)", re.I),
    re.compile(r"(\d{1,2})\+?\s*years?\s+(?:of\s+)?(?:experience|exp)", re.I),
    re.compile(r"(?:experience|background)[:\s]+(\d{1,2})\+?\s*years?", re.I),
    re.compile(r"\bacross\s+(\d{1,2})\+?\s*years?", re.I),
    re.compile(r"\b(\d{1,2})\s*years?\s*(?:\.|,|\s+of\b|\s+in\b|\s+as\b|\s+with\b|$)", re.I),
)

_EMAIL_RE = re.compile(r"^\s*[\w.+-]+@[\w.-]+\.\w+\s*$")
_URL_RE = re.compile(r"https?://|www\.", re.I)
_PHONE_RE = re.compile(r"^[\d\s\-+().]{10,}$")
_ROLE_LINE_HINT = re.compile(
    r"engineer|developer|scientist|architect|manager|analyst|designer|researcher|"
    r"\blead\b|staff|senior|principal|director|specialist|consultant|founder|cto|vp",
    re.I,
)


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    return re.sub(r"[ \t\r\f\v]+", " ", text)


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        chunks: list[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                chunks.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    chunks.append(" ".join(cells))
        return "\n".join(chunks)
    except Exception:
        return ""


def _is_pdf_magic(file_bytes: bytes) -> bool:
    return len(file_bytes) >= 4 and file_bytes[:4] == b"%PDF"


def _is_zip_magic(file_bytes: bytes) -> bool:
    return len(file_bytes) >= 4 and file_bytes[:2] == b"PK"


def _detect_format(file_bytes: bytes, mime: str) -> str | None:
    mime_l = (mime or "").lower()
    if "pdf" in mime_l or _is_pdf_magic(file_bytes):
        return "pdf"
    if "word" in mime_l or "officedocument.wordprocessingml" in mime_l:
        return "docx"
    if _is_zip_magic(file_bytes):
        return "docx"
    return None


def extract_plain_text(file_bytes: bytes, mime: str) -> str:
    """Best-effort plain text from PDF or DOCX."""
    fmt = _detect_format(file_bytes, mime)
    text = ""
    if fmt == "pdf":
        text = _extract_pdf_text(file_bytes)
    elif fmt == "docx":
        text = _extract_docx_text(file_bytes)
    else:
        if _is_pdf_magic(file_bytes):
            text = _extract_pdf_text(file_bytes)
        elif _is_zip_magic(file_bytes):
            text = _extract_docx_text(file_bytes)

    return _normalize_whitespace(text)


def _guess_years_experience(text: str) -> float:
    candidates: list[float] = []
    lower = text.lower()
    for pat in _YEARS_PATTERNS:
        for m in pat.finditer(lower):
            try:
                y = float(m.group(1))
                if 0 < y <= 45:
                    candidates.append(y)
            except (IndexError, ValueError):
                continue
    return max(candidates) if candidates else 0.0


def _skill_mentioned(lower: str, phrase: str) -> bool:
    pl = phrase.lower()
    if re.search(r"[^a-z0-9\s]", pl):
        return pl in lower
    if len(pl) <= 4:
        return re.search(rf"(?<![a-z0-9]){re.escape(pl)}(?![a-z0-9])", lower) is not None
    return pl in lower


def _guess_skills(text: str) -> list[str]:
    lower = text.lower()
    found: list[str] = []
    seen: set[str] = set()
    for phrase in sorted(_SKILL_PHRASES, key=len, reverse=True):
        pl = phrase.lower()
        if pl in seen:
            continue
        if _skill_mentioned(lower, phrase):
            seen.add(pl)
            found.append(phrase)
    return found[:25]


def _guess_headline(lines: list[str]) -> str:
    candidates: list[str] = []
    for line in lines[:25]:
        s = line.strip()
        if len(s) < 3 or len(s) > 140:
            continue
        if _EMAIL_RE.match(s) or _URL_RE.search(s) or _PHONE_RE.match(s):
            continue
        if re.match(r"^(resume|cv|curriculum vitae)\s*$", s, re.I):
            continue
        if re.match(r"^(education|experience|skills|summary|objective|work history)\s*$", s, re.I):
            continue
        if _ROLE_LINE_HINT.search(s):
            return s[:200]
        candidates.append(s[:200])
    return candidates[0] if candidates else ""


def _guess_name(lines: list[str]) -> str:
    """Rough heuristic: first short line that looks like a person's name."""
    for line in lines[:8]:
        s = line.strip()
        if not s or len(s) > 60:
            continue
        if _EMAIL_RE.match(s) or _URL_RE.search(s):
            continue
        if not re.match(
            r"^[A-Za-zÀ-ÖØ-öø-ÿ][A-Za-zÀ-ÖØ-öø-ÿ'.-]*(?:\s+[A-Za-zÀ-ÖØ-öø-ÿ'.-]+)+$",
            s,
        ):
            continue
        parts = s.split()
        if 2 <= len(parts) <= 5:
            return s
    return ""


def heuristic_profile_from_text(raw_text: str) -> dict[str, Any]:
    """Map extracted résumé text into ParsedProfile-shaped dict."""
    text = raw_text.strip()
    if not text:
        return {}

    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    skills = _guess_skills(text)
    years = _guess_years_experience(text)
    headline = _guess_headline(lines)
    name = _guess_name(lines)

    summary_cap = 2000
    summary = text[:summary_cap] if len(text) > summary_cap else text
    if len(text) > summary_cap:
        summary += "…"

    gaps: list[str] = []
    if not skills:
        gaps.append("No common technical keywords detected — add an explicit Skills section for better matching.")
    if years <= 0 and re.search(r"\b20\d{2}\b", text):
        gaps.append("Years of experience not inferred; consider stating tenure explicitly.")

    return {
        "name": name,
        "headline": headline,
        "experience_years": years,
        "skills": skills,
        "top_skills": skills[:10],
        "archetypes": [],
        "gaps": gaps,
        "summary": summary,
    }


async def parse_resume(file_bytes: bytes, mime: str) -> dict:
    """Extract structured profile data from résumé bytes (PDF or DOCX)."""
    text = extract_plain_text(file_bytes, mime)
    if not text.strip():
        return {
            "name": "",
            "headline": "",
            "experience_years": 0,
            "skills": [],
            "top_skills": [],
            "archetypes": [],
            "gaps": [
                "Could not extract text from this file. Try PDF text-based export or DOCX instead of a scanned image PDF."
            ],
            "summary": "No readable text found in the upload.",
        }

    return heuristic_profile_from_text(text)
